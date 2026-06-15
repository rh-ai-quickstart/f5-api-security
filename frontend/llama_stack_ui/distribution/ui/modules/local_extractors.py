# Copyright (c) Meta Platforms, Inc. and affiliates.
# All rights reserved.
#
# This source code is licensed under the terms described in the LICENSE file in
# the root directory of this source tree.

import io
import logging
import os

from docx import Document

logger = logging.getLogger(__name__)

LOCAL_SUPPORTED_EXTENSIONS = [".docx"]
PROVIDER_SUPPORTED_EXTENSIONS = [".txt", ".pdf", ".md"]


def extract_text_from_docx(file) -> str:
    """Extract all text content from a .docx file.

    Reads paragraph text and table cell text from the document.

    Args:
        file: File-like object containing .docx data

    Returns:
        str: Extracted text with paragraphs separated by newlines
    """
    doc = Document(file)
    parts = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def extract_text(file, filename: str) -> str:
    """Extract text from a locally supported file type.

    Routes to the appropriate extractor based on file extension.

    Args:
        file: File-like object with document data
        filename: Original filename used to determine the file type

    Returns:
        str: Extracted plain text content

    Raises:
        ValueError: If the file extension is not locally supported
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".docx":
        return extract_text_from_docx(file)
    else:
        raise ValueError(f"Unsupported file type for local extraction: {ext}")


def create_text_file_from_extracted_content(
    content: str, original_filename: str
) -> io.BytesIO:
    """Wrap extracted text as an in-memory .txt file for the Llama Stack API.

    Creates a BytesIO object with .name and .size attributes so it can be
    passed directly to the files.create API endpoint.

    Args:
        content: Extracted plain text to wrap
        original_filename: Original filename; the stem is reused with a .txt extension

    Returns:
        io.BytesIO: In-memory text file ready for upload
    """
    text_bytes = content.encode("utf-8")
    text_file = io.BytesIO(text_bytes)
    stem = os.path.splitext(original_filename)[0]
    text_file.name = f"{stem}.txt"
    text_file.size = len(text_bytes)
    return text_file
